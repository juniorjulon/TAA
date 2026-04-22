"""
generate_methodology_doc.py
===========================
Generates `docs/TAA_Methodology.docx` — the single, consolidated reference that
supersedes:
  - CLAUDE.md (system/implementation guide)
  - docs/TAA Signal Generation v1.0.md (methodology v1)
  - docs/TAA_System_Guide.md (system guide)
  - docs/TAA_Signal_Methodology.html (signal catalogue)

The document is a rigorous description of how TAA signals are generated,
aggregated, and mapped to asset classes — pulled directly from the current
authoritative source: `config/taa_config.xlsx`. Running this script always
produces a doc that matches the current Excel (so the doc never drifts from
the pipeline).
"""

from __future__ import annotations
import os
from collections import defaultdict
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
XLSX = os.path.join(ROOT, "config", "taa_config.xlsx")
OUT  = os.path.join(ROOT, "docs", "TAA_Methodology.docx")

PILLAR_ORDER = ("F", "M", "S", "V")
PILLAR_NAME  = {"F": "Fundamentals", "M": "Momentum", "S": "Sentiment", "V": "Valuation"}
PILLAR_COLOR = {"F": "14B8A6", "M": "F59E0B", "S": "A855F7", "V": "3A7BD5"}

BRAND  = RGBColor(0xC4, 0x12, 0x30)
INK    = RGBColor(0x0B, 0x12, 0x20)
SUB    = RGBColor(0x64, 0x74, 0x8B)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _read(ws):
    headers = [c.value for c in ws[1]]
    rows = []
    for r in ws.iter_rows(min_row=2, values_only=True):
        if all(v is None for v in r):
            continue
        rows.append(dict(zip(headers, r)))
    return rows


def load_config():
    wb = load_workbook(XLSX, data_only=True)
    return {
        "asset_classes":  _read(wb["AssetClasses"]),
        "data_series":    {s["series_id"]: s for s in _read(wb["DataSeries"])},
        "pillar_weights": {w["ac_id"]: w for w in _read(wb["PillarWeights"])},
        "pillar_notes":   {(n["ac_id"], n["pillar"]): n["note"] for n in _read(wb["PillarNotes"])},
        "mapping":        _read(wb["SignalMapping"]),
    }


def _shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def H1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BRAND
    return p


def H2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = INK
    return p


def H3(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = color or INK
    return p


def P(doc, text, italic=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def mono_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = INK
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Doc sections
# ─────────────────────────────────────────────────────────────────────────────
def section_cover(doc, cfg):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TAA Signal Generation — Methodology")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = BRAND

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = s.add_run("Consolidated reference for the Tactical Asset Allocation signal system")
    rr.italic = True
    rr.font.size = Pt(11)
    rr.font.color.rgb = SUB

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    import datetime as _dt
    dr = d.add_run(f"Generated from config/taa_config.xlsx on {_dt.date.today().isoformat()}")
    dr.font.size = Pt(9)
    dr.font.color.rgb = SUB
    doc.add_paragraph()


def section_executive_summary(doc, cfg):
    H1(doc, "Executive Summary")
    P(doc,
      "This system produces tactical tilts around strategic benchmarks for an "
      f"insurance portfolio across {len(cfg['asset_classes'])} asset classes. "
      "Each asset class is scored along four pillars — Fundamentals (F), "
      "Momentum (M), Sentiment (S), and Valuation (V) — whose z-scores are "
      "weighted into a composite score that is mapped to a conviction-based "
      "tilt capped at ±5% per asset class with a 50–150bp tracking-error budget."
    )
    P(doc,
      "The authoritative configuration lives in config/taa_config.xlsx. That "
      "workbook drives (a) the methodology blueprint in index.html, (b) the "
      "parameter block of src/config.py, and (c) this document. Running "
      "`python src/build_dashboard.py` pushes Excel changes into the HTML and "
      "Python layers; running `python src/generate_methodology_doc.py` "
      "refreshes this Word file.")

    H2(doc, "Pipeline in one picture")
    mono_block(doc,
        "config/taa_config.xlsx  (user-owned source of truth)\n"
        "          │\n"
        "          ├─▶ src/build_dashboard.py\n"
        "          │       ├─▶ index.html  (FI/EQ blueprints, SIG_MATRIX, AC meta)\n"
        "          │       └─▶ src/config.py  (ASSET_CLASSES, PILLAR_WEIGHTS, MAX_TILT_PCT)\n"
        "          │\n"
        "          └─▶ src/generate_methodology_doc.py\n"
        "                  └─▶ docs/TAA_Methodology.docx  (this file)"
    )


def section_universe(doc, cfg):
    H1(doc, "Part 1 — Asset Class Universe")
    P(doc, "All twelve asset classes carried by the system, grouped into Fixed "
           "Income (FI) and Equity (EQ).")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["ID", "Full Label", "Short", "Group", "Max tilt"]):
        hdr[i].text = h
        _shading(hdr[i], "1E2E47")
        for r in hdr[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
    for a in cfg["asset_classes"]:
        row = table.add_row().cells
        row[0].text = a["ac_id"]
        row[1].text = a["full_label"]
        row[2].text = a["short_label"]
        row[3].text = a["group"]
        row[4].text = f"±{float(a['max_tilt_pct']):.1f}%"


def section_data_series(doc, cfg):
    H1(doc, "Part 2 — Data Series Catalogue")
    P(doc, "Every raw signal tracked by the system, organised by pillar. Each "
           "series is mapped to one or more asset classes in Part 3.")
    by_pillar = defaultdict(list)
    for sid, s in cfg["data_series"].items():
        by_pillar[s["pillar"]].append(s)
    for p in PILLAR_ORDER:
        H2(doc, f"Pillar {p} — {PILLAR_NAME[p]}")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Series ID", "Signal", "Ticker", "Source", "Transformation"]):
            hdr[i].text = h
            _shading(hdr[i], PILLAR_COLOR[p])
            for r in hdr[i].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.bold = True
        for s in by_pillar[p]:
            row = table.add_row().cells
            row[0].text = str(s["series_id"])
            row[1].text = str(s["signal_name"])
            row[2].text = str(s["ticker"] or "")
            row[3].text = str(s["source"] or "")
            row[4].text = str(s["transformation"] or "")


def section_aggregation(doc, cfg):
    H1(doc, "Part 3 — Aggregation Pipeline")
    P(doc, "Raw series become a single composite score per asset class via five "
           "deterministic steps.")

    H2(doc, "Step 1 — Normalize each raw series to a z-score")
    bullet(doc, "EWMA z-score (default for daily signals, span = 252×3 days ≈ 3Y half-life).")
    bullet(doc, "Rolling z-score for slow-moving valuation signals (P/E, ERP, real yield) with 10Y window.")
    bullet(doc, "Percentile rank (5Y or 10Y) for non-normal distributions: VIX, OAS, P/E.")
    bullet(doc, "All z-scores winsorised to ±3σ.")

    H2(doc, "Step 2 — Apply sign conventions")
    bullet(doc, "Fixed-income duration: growth signals (PMI/GDP/CESI) are inverted (−1×).")
    bullet(doc, "Credit: spread tightening is positive → OAS/CDS momentum multiplied by −1.")
    bullet(doc, "Yields: falling yield = rising bond price → yield momentum multiplied by −1.")
    bullet(doc, "VIX / MOVE at extremes apply contrarian scoring (sign C): pctile > 0.85 → +1.5; pctile < 0.15 → −1.5.")
    bullet(doc, "CESI at percentile > 0.85 or < 0.15 is inverted (mean-reverting).")

    H2(doc, "Step 3 — Aggregate within each pillar")
    P(doc, "Signals are combined inside each pillar using the weights specified "
           "in SignalMapping. The result is re-standardised to restore unit variance.")
    mono_block(doc,
        "pillar_F_us = Σ w_i × z_i   over Fundamentals signals mapped to us_equity\n"
        "Z_F_us      = rolling_zscore(pillar_F_us, 252).clip(-3, 3)"
    )

    H2(doc, "Step 4 — Combine pillars into a composite")
    P(doc, "Per-asset-class pillar weights (from PillarWeights) feed the composite.")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Asset Class", "F", "M", "S", "V", "Σ"]):
        hdr[i].text = h
        _shading(hdr[i], "1E2E47")
        for r in hdr[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
    for a in cfg["asset_classes"]:
        w = cfg["pillar_weights"][a["ac_id"]]
        tot = sum(float(w[p]) for p in PILLAR_ORDER)
        row = table.add_row().cells
        row[0].text = a["full_label"]
        for i, p in enumerate(PILLAR_ORDER, start=1):
            row[i].text = f"{float(w[p]) * 100:.0f}%"
        row[5].text = f"{tot * 100:.0f}%"

    H2(doc, "Step 5 — Apply the pillar-agreement multiplier")
    P(doc, "Conviction quality filter: count how many of F/M/S/V share the sign "
           "of the composite (ignoring pillars with |z| < 0.25).")
    bullet(doc, "4 of 4 agree → multiplier 1.00 — full conviction.")
    bullet(doc, "3 of 4 agree → multiplier 0.80.")
    bullet(doc, "2 of 4 agree → multiplier 0.50.")
    bullet(doc, "1 or 0 agree → multiplier 0.00 — the pillars disagree too much; tilt is neutralised.")


def section_abs_vs_rel(doc, cfg):
    H1(doc, "Part 4 — Absolute vs Relative Views, Tilt Sizing")
    P(doc, "The composite is blended 35% absolute / 65% relative before sizing the "
           "tilt. Absolute view asks: \"Is this AC attractive vs its own history?\" "
           "Relative view asks: \"Which AC do I prefer over the others?\" The two "
           "together let the system express \"OW equity / UW FI\" calls naturally "
           "even when both absolute scores are positive.")
    mono_block(doc,
        "Z_rel    = (Z_composite - mean(Z_composite across AC)) / std(...)\n"
        "Z_final  = 0.35 × Z_composite  +  0.65 × Z_rel\n"
        "tilt_ac  = clip(Z_final × max_tilt_ac × conviction_mult, -max_tilt_ac, +max_tilt_ac)"
    )
    H2(doc, "Conviction mapping (absolute view)")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Z composite", "Label", "Tilt vs SAA"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        _shading(cell, "1E2E47")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
    for z, lbl, t in [
        ("> +1.50",         "HIGH OVERWEIGHT",   "+3.0% to +5.0%"),
        ("+0.75 to +1.50",  "MEDIUM OVERWEIGHT", "+1.5% to +3.0%"),
        ("-0.75 to +0.75",  "NEUTRAL",           "0%"),
        ("-1.50 to -0.75",  "MEDIUM UNDERWEIGHT","-1.5% to -3.0%"),
        ("< -1.50",         "HIGH UNDERWEIGHT",  "-3.0% to -5.0%"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = z, lbl, t


def section_crisis(doc, cfg):
    H1(doc, "Part 5 — Crisis Override")
    P(doc, "Hard rule on top of the conviction mapping: when VIX percentile > 0.80 "
           "AND MOVE percentile > 0.80 simultaneously, all tilts are forced to zero "
           "until both indicators fall below the 70th percentile. This dominates "
           "any individual composite signal.")


def section_blueprints(doc, cfg):
    H1(doc, "Appendix — Signal Blueprints per Asset Class")
    P(doc, "For every asset class, the complete signal mapping by pillar: which "
           "series, sign convention, weight inside the pillar, and a note "
           "explaining the economic intuition.")

    mapping_by_ac_pil = defaultdict(list)
    for m in cfg["mapping"]:
        mapping_by_ac_pil[(m["ac_id"], m["pillar"])].append(m)

    for a in cfg["asset_classes"]:
        ac_id = a["ac_id"]
        H2(doc, f"{a['full_label']}")
        P(doc, f"{a['sub_description']} — benchmark: {a['benchmark']} · max tilt ±{float(a['max_tilt_pct']):.1f}%",
          italic=True, size=9)
        w = cfg["pillar_weights"][ac_id]
        for p in PILLAR_ORDER:
            title = f"{PILLAR_NAME[p]}  ({int(round(float(w[p]) * 100))}%)"
            H3(doc, title, color=RGBColor(int(PILLAR_COLOR[p][0:2], 16),
                                          int(PILLAR_COLOR[p][2:4], 16),
                                          int(PILLAR_COLOR[p][4:6], 16)))
            note = cfg["pillar_notes"].get((ac_id, p))
            if note:
                P(doc, note, italic=True, size=9)
            sigs = mapping_by_ac_pil.get((ac_id, p), [])
            if not sigs:
                P(doc, "(no signals mapped)", italic=True, size=9)
                continue
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(["Signal", "Description", "Sign", "Weight"]):
                cell = table.rows[0].cells[i]
                cell.text = h
                _shading(cell, PILLAR_COLOR[p])
                for r in cell.paragraphs[0].runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.bold = True
            for m in sigs:
                s_meta = cfg["data_series"].get(m["series_id"], {})
                desc = m.get("description_override") or s_meta.get("notes") or ""
                wt = m.get("weight_in_pillar") or ""
                if isinstance(wt, (int, float)):
                    wt = f"{int(round(float(wt) * 100))}%" if float(wt) <= 1 else f"{int(round(float(wt)))}%"
                row = table.add_row().cells
                row[0].text = s_meta.get("signal_name", m["series_id"]) or ""
                row[1].text = str(desc)
                row[2].text = str(m.get("sign") or "—")
                row[3].text = str(wt)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
def main():
    if not os.path.isfile(XLSX):
        raise SystemExit(f"Config Excel not found: {XLSX}")
    cfg = load_config()

    doc = Document()
    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    section_cover(doc, cfg)
    section_executive_summary(doc, cfg)
    section_universe(doc, cfg)
    section_data_series(doc, cfg)
    section_aggregation(doc, cfg)
    section_abs_vs_rel(doc, cfg)
    section_crisis(doc, cfg)
    section_blueprints(doc, cfg)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"[OK] wrote {OUT}")


if __name__ == "__main__":
    main()
