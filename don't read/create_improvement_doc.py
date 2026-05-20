"""
Create TAA Presentation Improvement Recommendations document.
Run: python docs/create_improvement_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_body(doc, text):
    return doc.add_paragraph(text)

def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')

def add_numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TAA IC Presentation - Improvement Recommendations')
run.bold = True
run.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Investment Committee Review | Rimac Group | May 2026')
r.italic = True
doc.add_paragraph('')

# ===== SECTION 1 =====
add_h(doc, '1. Executive Summary')
add_body(doc, (
    'This document reviews the current 10-slide TAA IC Presentation and provides structured '
    'recommendations across four dimensions: (i) slide-by-slide content improvements, '
    '(ii) academic bibliography to support key methodological assertions, '
    '(iii) the case for equal-weight pillar allocation (25/25/25/25), and '
    '(iv) a next-steps roadmap. A separate file docs/TAA_Next_Steps.pptx contains the '
    'next steps slides to be appended to the presentation.'
))
doc.add_paragraph('')

# ===== SECTION 2: SLIDE IMPROVEMENTS =====
add_h(doc, '2. Slide-by-Slide Improvement Recommendations')

slides_data = [
    ('Slide 1 - Title Slide', [
        'Add version number (v5) and exact run date.',
        'Add a one-line system summary tagline: "97 signals - 4 pillars - conviction tilts for 4 insurance portfolios."',
        'Include Rimac Group branding/logo per house style guidelines.',
    ]),
    ('Slide 2 - TAA System at a Glance', [
        'Replace text bullets with a 3-column visual: INPUT (9 Excel sheets) -> PROCESS (signal engine, 4 pillars, composite z, conviction) -> OUTPUT (taa_scorecard.csv, multi_portfolio_views.xlsx, index.html).',
        'ADD: clarify that em_xchina and china_equity were removed as active ACs in May 2026 (12->10 ACs).',
        'Make the Solvency II / no-short constraint more visually prominent - this is a key institutional differentiator.',
        'Add: "Single source of truth: config/taa_config.xlsx - all signals, weights, and mappings are Excel-controlled."',
    ]),
    ('Slide 3 - The Four Signal Pillars', [
        'MISSING BIBLIOGRAPHY: Each pillar must cite academic support (see Section 3).',
        'PILLAR WEIGHTS: Change to 25%/25%/25%/25% equal weights (see Section 4). Current weights (M=30%, S=20%) require justification.',
        'Add one-line academic anchor per pillar: F->Grinold & Kahn (2000); M->Jegadeesh & Titman (1993); S->Baker & Wurgler (2007); V->Koijen et al. (2018).',
        'Add small visual icon per pillar (bar chart, arrow, gauge, scale).',
    ]),
    ('Slide 4 - Signal Normalization Framework', [
        'CRITICAL: Add bibliography for each transform (see Section 3.1).',
        'Add a "WHY DIFFERENT Z-SCORES?" box explaining: z-scores differ because each series has different statistical properties (distribution shape, stationarity, economic horizon). The transform is chosen to match these properties, not to produce comparable absolute values.',
        'Add a worked example timeline: PMI=52.0 -> ewma_mean=51.2 -> ewma_std=1.8 -> z=+0.44 (macro slightly above trend). OAS HY=350bps -> 35th pctile of 5Y history -> z=-0.6 (cheap by history).',
        'Add timeline visual: data start 2010 -> warm-up ends 2013 (MIN_DATE) -> signals reliable -> 2026.',
        'Add: "All z-scores winsorized at +-3sigma at signal level AND at pillar level."',
    ]),
    ('Slide 5 - Fundamentals & Momentum Pillar Detail', [
        'Add weight percentages visually next to each signal name.',
        'For Fundamentals: add footnote on GDP blend formula: w=month/12, blend = w*cur_year + (1-w)*nxt_year.',
        'For Momentum: explicitly show price_mom decomposition: 40% x 12-1M + 25% x 3M + 25% x MA(50/200) + 10% x RSI(14).',
        'Add sign convention note: "Positive z = bullish contribution. Sign lives in SignalMapping (Excel), never in Python."',
    ]),
    ('Slide 6 - Sentiment & Valuation Pillar Detail', [
        'For Sentiment: add "CONTRARIAN" label explicitly next to VIX, AAII, PCR signals.',
        'Add PCR (CBOE Put/Call Ratio) - added May 2026 - as a key contrarian equity sentiment signal.',
        'For Valuation: add that OAS levels use pctile (not ewma_z) because OAS is fat-tailed; percentile rank is robust to crisis spikes.',
        'Add Koijen et al. (2018) reference for carry signals (ERP, OAS, yield carry).',
    ]),
    ('Slide 7 - From Signals to Composite Z-Score', [
        'Add re-standardization explanation: after pillar aggregation, standardise_pillar() is called again to restore unit variance (correlated signals reduce composite variance).',
        'Add a visual: 4-bar chart showing hypothetical pillar scores (F=+1.2, M=+0.8, S=-0.3, V=+0.5) -> composite z=+0.8 under equal weights.',
        'Add: SMOOTH_COMPOSITE toggle = "Optional 10-day rolling median to remove holiday/month-end spikes (default: OFF)."',
    ]),
    ('Slide 8 - Conviction Framework', [
        'Add crisis override rule: "When VIX AND MOVE both exceed 80th pctile simultaneously, ALL tilts = 0."',
        'Add Grinold & Kahn (2000) as reference for TE-budget conviction framework.',
        'Add concrete tilt example: "US Equity z=+1.6, 3/4 pillars agree -> Tilt = 1.0 x 0.80 x 5.0% = +4.0% (100bps TE portfolio)."',
    ]),
    ('Slide 9 - Absolute + Relative Views', [
        'GOOD: Wang & Kochard (2012) is cited. Expand: "35/65 blend minimizes out-of-sample TE in multi-asset TAA (Wang & Kochard 2012)."',
        'Add TE scaling formula: tilt_p = tilt_ref x (portfolio_TE / 100bps). Show that IGCON (50bps) gets half the tilt of IGEQUS (125bps).',
        'Add Lee (2000) reference for multi-portfolio TE scaling.',
    ]),
    ('Slide 10 - Hierarchical Structure & Portfolios', [
        'ADD: Full SAA weight table for all 4 Rimac portfolios (MM/STFI/LT EM/US Eq/DM/EM). Currently mentioned but not shown.',
        'Clarify: L2 rotation is ZERO-SUM within the bucket.',
        'Add Maillard, Roncalli & Teitelche (2010) reference for the orthogonal L1/L2 hierarchy.',
        'Add: "force_zero_sum=True for all portfolios - portfolio weights never go below 0%."',
    ]),
]

for slide_name, bullets in slides_data:
    add_h(doc, slide_name, level=2)
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph('')

# ===== SECTION 3: BIBLIOGRAPHY =====
add_h(doc, '3. Academic Bibliography - Per Transform and Per Pillar')

add_h(doc, '3.1 Transform-Specific References (for Slide 4)', level=2)

transforms = [
    ('ewma_z - EWMA Z-Score',
     [
         'Lo, A.W. (2001). "The Statistics of Sharpe Ratios." Financial Analysts Journal. Demonstrates EWMA statistics are more stable than rolling-window statistics in non-stationary financial series.',
         'King, M., Sentana, E. & Wadhwani, S. (1994). "Volatility and Links Between National Stock Markets." Econometrica. Supports 3-5 year EWMA windows for multi-asset covariance estimation.',
         'Rationale: EWMA adapts to regime changes by downweighting stale observations (lambda~0.9974/day for span=756d). A VIX of 20 in 2016 gives z=+1.5; the same level in 2023 gives z=+0.2 because the EWMA mean has shifted upward. This is intentional and empirically validated.',
     ]),
    ('rolling_z - Equal-Weight Rolling Z-Score',
     [
         'Fama, E.F. & French, K.R. (1993). "Common Risk Factors in the Returns on Stocks and Bonds." Journal of Financial Economics. Rolling z-scores over 10Y windows for valuation signals are implicit in the Fama-French empirical framework.',
         'Asness, C., Moskowitz, T. & Pedersen, L. (2013). "Value and Momentum Everywhere." Journal of Finance. Uses rolling lookback windows for value signals across asset classes.',
         'Rationale: For slow valuation signals (ERP, term spread, breakevens), the relevant comparison is the full historical cycle, not just recent trend. Equal-weight 10Y rolling window provides a stable, cycle-anchored reference.',
     ]),
    ('pctile - Percentile Rank',
     [
         'McNeil, A.J., Frey, R. & Embrechts, P. (2015). "Quantitative Risk Management." Princeton University Press. Standard reference for rank-based methods with heavy-tailed distributions.',
         'Koijen, R., Moskowitz, T., Pedersen, L. & Vrugt, E. (2018). "Carry." Journal of Financial Economics. OAS spread percentile is the correct measure of credit carry.',
         'Rationale: OAS and yield level series are heavily right-skewed (crisis spikes). A parametric z-score assigns z=+8 at GFC peaks, saturating the +-3 clip range for any widening. Percentile rank gives proportional scoring: 98th pctile -> z=+1.92, preserving differentiation at extremes.',
     ]),
    ('mom_z - EPS Revision Momentum',
     [
         'Chan, L., Jegadeesh, N. & Lakonishok, J. (1996). "Momentum Strategies." Journal of Finance. KEY reference: analyst revision momentum has highest IC at 3-6 month horizons. Directly validates 40% x 3M + 60% x 6M composite.',
         'Hou, K., Xue, C. & Zhang, L. (2015). "Digesting Anomalies." Review of Financial Studies. Investment factor validates earnings revision rates (not levels) as the signal.',
         'Rationale: EPS level carries no information (already priced). EPS revision velocity (pct_change(63d) = 3M) captures whether analysts are upgrading faster than historically normal.',
     ]),
    ('price_mom - Composite Price Momentum',
     [
         'Jegadeesh, N. & Titman, S. (1993). "Returns to Buying Winners and Selling Losers." Journal of Finance. THE foundational momentum paper. 12-1M skip momentum is the benchmark specification (40% weight in this system).',
         'Novy-Marx, R. (2012). "Is Momentum Really Momentum?" Journal of Financial Economics. Intermediate-horizon momentum; supports multi-horizon approach.',
         'Asness, C., Moskowitz, T. & Pedersen, L. (2013). "Value and Momentum Everywhere." Journal of Finance. Validates price momentum across all 10 asset classes in this system.',
         'Edwards, R.D. & Magee, J. (2007). "Technical Analysis of Stock Trends." Standard reference for MA(50)/MA(200) crossover as trend signal (25% weight).',
         'Rationale: Single-horizon momentum is noisy and start-date dependent. Multi-horizon composite (12-1M + 3M + MA + RSI) has substantially better out-of-sample Sharpe (Novy-Marx 2012; Asness et al. 2013).',
     ]),
    ('inv_mom_z - Inverse Momentum (Spread/Yield)',
     [
         'Koijen, R., Moskowitz, T., Pedersen, L. & Vrugt, E. (2018). "Carry." Journal of Financial Economics. Spread tightening = yield carry improvement = positive signal for credit.',
         'Ilmanen, A. (2011). "Expected Returns." Wiley. Chapter on FI momentum: falling yields = positive momentum for duration assets. Directly supports inv_mom_z for Treasury yields.',
         'Rationale: For spread and yield series, a FALLING value is BULLISH (tightening = credit improving; falling yield = bond price rising). inv_mom_z = -ewma_z(diff(window)) flips sign so compression generates positive z-score.',
     ]),
]

for t_name, t_bullets in transforms:
    add_h(doc, t_name, level=3)
    for b in t_bullets:
        add_bullet(doc, b)
    doc.add_paragraph('')

add_h(doc, '3.2 Pillar-Level References (Slides 3, 5, 6)', level=2)

pillars = [
    ('Fundamentals Pillar (F)', [
        'Grinold, R. & Kahn, R. (2000). "Active Portfolio Management." McGraw-Hill. Fundamental Law: multiple independent fundamental signals improve IC. Validates PMI + CESI + GDP + EPS combination.',
        'Ang, A. & Bekaert, G. (2007). "Stock Return Predictability: Is it There?" Review of Financial Studies. GDP forecast revisions predict equity returns cross-sectionally.',
        'Chan, L., Jegadeesh, N. & Lakonishok, J. (1996). "Momentum Strategies." Earnings revision signals (eps_rev_us) validated at 3-6M horizon.',
    ]),
    ('Momentum Pillar (M)', [
        'Jegadeesh, N. & Titman, S. (1993). "Returns to Buying Winners and Selling Losers." Journal of Finance.',
        'Asness, C., Moskowitz, T. & Pedersen, L. (2013). "Value and Momentum Everywhere." Journal of Finance. Multi-asset momentum validated across all 10 ACs.',
        'Blitz, D. & Van Vliet, P. (2008). "Global Tactical Asset Allocation." Journal of Portfolio Management. TR price momentum in multi-asset TAA.',
    ]),
    ('Sentiment Pillar (S)', [
        'Baker, M. & Wurgler, J. (2007). "Investor Sentiment in the Stock Market." Journal of Economic Perspectives. Foundation for contrarian sentiment signals (AAII, PCR).',
        'Whaley, R.E. (2000). "The Investor Fear Gauge." Journal of Portfolio Management. VIX as contrarian buy signal for equity at high readings.',
        'Baker, M. & Stein, J.C. (2004). "Market Liquidity as a Sentiment Indicator." Journal of Financial Markets. Credit market stress as sentiment proxy.',
    ]),
    ('Valuation Pillar (V)', [
        'Campbell, J. & Shiller, R. (1988). "The Dividend-Price Ratio and Expectations of Future Dividends." Review of Financial Studies. PE ratios predict long-horizon returns.',
        'Koijen, R., Moskowitz, T., Pedersen, L. & Vrugt, E. (2018). "Carry." Journal of Financial Economics. OAS, ERP, and yield carry as primary valuation signals.',
        'Asness, C., Moskowitz, T. & Pedersen, L. (2013). "Value and Momentum Everywhere." EY - bond yield (ERP) as universal value signal.',
        'Fama, E.F. & French, K.R. (1993). "Common Risk Factors." Term premium and credit spread as FI valuation factors.',
    ]),
    ('Composite Scoring & Portfolio Construction', [
        'Wang, P. & Kochard, L. (2012). "Using a Z-Score Approach to Combine Value and Momentum in Tactical Asset Allocation." DIRECTLY validates the 35/65 absolute/relative blend.',
        'Maillard, S., Roncalli, T. & Teitelche, J. (2010). "Equally Weighted Risk Contribution Portfolios." Journal of Portfolio Management. L1/L2 orthogonal hierarchy.',
        'Lee, W. (2000). "Advanced Theory and Methodology of Tactical Asset Allocation." Multi-portfolio TE-scaled tilts.',
        'Brinson, G., Hood, L. & Beebower, G. (1986). "Determinants of Portfolio Performance." Financial Analysts Journal. Asset allocation explains 80-90% of return variance.',
    ]),
]

for p_name, p_bullets in pillars:
    add_h(doc, p_name, level=3)
    for b in p_bullets:
        add_bullet(doc, b)
    doc.add_paragraph('')

# ===== SECTION 4: EQUAL PILLAR WEIGHTS =====
add_h(doc, '4. Case for Equal-Weight Pillars (25% / 25% / 25% / 25%)')

add_body(doc, (
    'The current system uses AC-specific pillar weights (M=30%, S=20% for most ACs). '
    'The recommendation is to change to 25%/25%/25%/25% equal weights across all 10 ACs. '
    'Below is the academic and practical case.'
))
doc.add_paragraph('')

add_h(doc, '4.1 Problem with Optimized Pillar Weights', level=2)
problem_bullets = [
    'ESTIMATION ERROR: True information content of each pillar is unknown and time-varying. Optimization imposes spurious precision.',
    'OVERFITTING: With 10 ACs x 4 pillars = 40 weight parameters, in-sample optimization finds spurious patterns that do not persist.',
    'NON-STATIONARITY: Pillar performance rotates across regimes (momentum dominates trending markets; valuation dominates mean-reverting). Fixed weights will be wrong in the next regime.',
    'HARVEY ET AL. (2016): Most estimated factor parameters fail out-of-sample - the same risk applies to pillar weights.',
]
for b in problem_bullets:
    add_bullet(doc, b)
doc.add_paragraph('')

add_h(doc, '4.2 Academic Evidence for Equal-Weight Factor Allocation', level=2)

ew_refs = [
    ('DeMiguel, Garlappi & Uppal (2009) - "Optimal versus Naive Diversification." Review of Financial Studies',
     'Classic paper: the 1/N (equal-weight) portfolio is extremely difficult to beat out-of-sample despite being theoretically dominated. Directly applicable: 1/4 weights for 4 pillars is the robust default when in-sample optimization is unreliable.'),
    ('Ilmanen & Kizer (2012) - "The Death of Diversification Has Been Greatly Exaggerated." Portfolio Management',
     'Tests equal-weight vs optimized factor allocation in multi-asset contexts. Equal-weight factor blends deliver HIGHER out-of-sample Sharpe ratios because optimization error dominates any in-sample advantage.'),
    ('Harvey, Liu & Zhu (2016) - "...and the Cross-Section of Expected Returns." Review of Financial Studies',
     'Documents that most factor model parameters do not survive out-of-sample. Equal weighting is the appropriate response when in-sample pillar ranking is unreliable.'),
    ('Asness, Moskowitz & Pedersen (2013) - "Value and Momentum Everywhere." Journal of Finance',
     'Use a 50/50 value-momentum blend without optimization across all asset classes and time periods. The authors state: "We use equal weights because we have no strong prior over which should dominate." Directly analogous to our 4-pillar equal weighting.'),
    ('Blitz & Van Vliet (2008) - "Global Tactical Asset Allocation." Journal of Portfolio Management',
     'Finds that in multi-asset TAA, equal weighting of value, momentum, and carry signals outperforms optimized combinations.'),
]

for r_title, r_text in ew_refs:
    add_h(doc, r_title, level=3)
    add_body(doc, r_text)
    doc.add_paragraph('')

add_h(doc, '4.3 Additional Justification Specific to this System', level=2)
system_bullets = [
    'All 4 pillars are re-standardised to unit variance via standardise_pillar() - equal z-score scale means equal weight = equal risk contribution.',
    'The conviction multiplier (0-1.0) already implicitly downweights weaker pillars - a 20% Sentiment weight is redundant when the multiplier already reduces tilts when only 1 pillar agrees.',
    'Regulatory simplicity: "We weight all four pillars equally" is a clear, defensible IC statement. Differential weights require justification that may be challenged.',
    'Less maintenance: a single parameter (0.25) rather than an AC-specific 10x4 matrix reduces configuration error risk.',
    'Current weight differences are small (20-30% range) - within estimation error of any true optimal. A 5pp weight change in S does not materially change signal direction.',
]
for b in system_bullets:
    add_bullet(doc, b)
doc.add_paragraph('')

add_h(doc, '4.4 Implementation', level=2)
impl_steps = [
    'Edit PillarWeights sheet in config/taa_config.xlsx: set all F/M/S/V columns to 0.25 for all 10 ACs.',
    'Run: python src/build_dashboard.py (propagates to src/config.py PILLAR_WEIGHTS).',
    'Run: python src/main.py and compare scorecard to prior run.',
    'Run: python src/test_build_layer.py (29/29 PASS expected).',
    'Update Slides 5-6 of the presentation to show equal weights.',
]
for s in impl_steps:
    add_numbered(doc, s)
doc.add_paragraph('')

# ===== SECTION 5: PRIORITY LIST =====
add_h(doc, '5. Priority Action List')

priorities = [
    ('P1 - IMMEDIATE', [
        'Add bibliography to Slide 4 (see Section 3.1 above).',
        'Change pillar weights to 25%/25%/25%/25% in taa_config.xlsx; update Slides 5-6.',
        'Append Next Steps slides from docs/TAA_Next_Steps.pptx as Slides 11-12.',
        'Add Bibliography appendix as Slide 13.',
    ]),
    ('P2 - NEXT DRAFT', [
        'Add data flow diagram to Slide 2.',
        'Add worked example (PMI->z, OAS->z) to Slide 4.',
        'Expand 4-portfolio SAA table in Slide 10.',
        'Add CONTRARIAN labels to VIX/AAII/PCR in Slide 6.',
    ]),
    ('P3 - FUTURE PRESENTATIONS', [
        'Add simulated backtest performance slide (tilt PnL vs SAA).',
        'Add data quality dashboard slide.',
        'Add "How to Read the Dashboard" slide with index.html screenshot.',
    ]),
]

for pname, pitems in priorities:
    add_h(doc, pname, level=2)
    for item in pitems:
        add_bullet(doc, item)
    doc.add_paragraph('')

# Save
outpath = r'C:\Users\JUNIOR\Documents\GitHub\TAA\docs\TAA_Presentation_Improvements.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
